import logging
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class DOCXParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            import docx
            doc = docx.Document(source)
            text_blocks = [para.text for para in doc.paragraphs if para.text.strip()]
            
            tables = []
            for table in doc.tables:
                grid = []
                for row in table.rows:
                    grid.append([cell.text.strip() for cell in row.cells])
                tables.append(grid)
                
            return self._standard_response(
                parser_name="docx",
                success=True,
                text="\n".join(text_blocks),
                tables=tables,
                confidence=0.95
            )
        except ImportError:
            logger.error("python-docx not installed.")
            return self._standard_response(parser_name="docx", success=False, confidence=0.0)
        except Exception as e:
            logger.error(f"DOCXParser failed: {e}")
            return self._standard_response(parser_name="docx", success=False, confidence=0.0)
