import os

base_path = "/mnt/c/Users/khand/GovTrackAI/parsers"

# 1. base_parser.py
base_parser_code = """from abc import ABC, abstractmethod
import logging

logger = logging.getLogger(__name__)

class BaseParser(ABC):
    @abstractmethod
    def parse(self, source, **kwargs) -> dict:
        \"\"\"
        Parse the source (file path or content) and return a standardized dictionary.
        \"\"\"
        raise NotImplementedError("Subclasses must implement parse()")

    def _standard_response(self, parser_name: str, success: bool = True, text: str = "", metadata: dict = None, structured_data: dict = None, tables: list = None, attachments: list = None, confidence: float = 1.0) -> dict:
        return {
            "success": success,
            "parser": parser_name,
            "metadata": metadata or {},
            "text": text,
            "structured_data": structured_data or {},
            "tables": tables or [],
            "attachments": attachments or [],
            "confidence": confidence
        }
"""
with open(os.path.join(base_path, "base_parser.py"), "w") as f:
    f.write(base_parser_code)


# 2. factory.py
factory_code = """import os
import logging
from parsers.base_parser import BaseParser

# Lazy loading to prevent circular imports
logger = logging.getLogger(__name__)

class ParserFactory:
    @staticmethod
    def get_parser(file_path_or_url: str) -> BaseParser:
        \"\"\"Automatically detect file type and return the correct parser instance.\"\"\"
        ext = file_path_or_url.lower().split('.')[-1].split('?')[0]
        
        if ext == 'pdf':
            from parsers.pdf_parser import PDFParser
            return PDFParser()
        elif ext in ['html', 'htm', 'php', 'aspx']:
            from parsers.html_parser import HTMLParser
            return HTMLParser()
        elif ext == 'json':
            from parsers.json_parser import JSONAPIParser
            return JSONAPIParser()
        elif ext == 'xml' or ext == 'rss':
            from parsers.rss_parser import RSSParser
            return RSSParser()
        elif ext == 'zip':
            from parsers.zip_parser import ZIPParser
            return ZIPParser()
        elif ext == 'docx':
            from parsers.docx_parser import DOCXParser
            return DOCXParser()
        elif ext == 'txt':
            from parsers.txt_parser import TXTParser
            return TXTParser()
        
        logger.warning(f"Unsupported extension '{ext}'. Defaulting to TXTParser.")
        from parsers.txt_parser import TXTParser
        return TXTParser()
"""
with open(os.path.join(base_path, "factory.py"), "w") as f:
    f.write(factory_code)


# 3. json_parser.py
json_parser_code = """import json
import logging
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class JSONAPIParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            with open(source, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            text_rep = json.dumps(data, indent=2)
            return self._standard_response(
                parser_name="json",
                success=True,
                text=text_rep,
                structured_data=data if isinstance(data, dict) else {"items": data},
                confidence=1.0
            )
        except Exception as e:
            logger.error(f"JSONAPIParser failed: {e}")
            return self._standard_response(parser_name="json", success=False, confidence=0.0)
"""
with open(os.path.join(base_path, "json_parser.py"), "w") as f:
    f.write(json_parser_code)


# 4. rss_parser.py
rss_parser_code = """import logging
from parsers.base_parser import BaseParser
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class RSSParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            tree = ET.parse(source)
            root = tree.getroot()
            
            items = []
            text_blocks = []
            
            for item in root.findall('.//item'):
                title = item.find('title')
                desc = item.find('description')
                link = item.find('link')
                
                t = title.text if title is not None else ""
                d = desc.text if desc is not None else ""
                l = link.text if link is not None else ""
                
                items.append({"title": t, "description": d, "link": l})
                text_blocks.append(f"{t}\\n{d}")
                
            return self._standard_response(
                parser_name="rss",
                success=True,
                text="\\n\\n".join(text_blocks),
                structured_data={"items": items},
                confidence=0.9
            )
        except Exception as e:
            logger.error(f"RSSParser failed: {e}")
            return self._standard_response(parser_name="rss", success=False, confidence=0.0)
"""
with open(os.path.join(base_path, "rss_parser.py"), "w") as f:
    f.write(rss_parser_code)


# 5. zip_parser.py
zip_parser_code = """import zipfile
import logging
import os
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class ZIPParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            attachments = []
            text_blocks = []
            
            with zipfile.ZipFile(source, 'r') as z:
                for file_info in z.infolist():
                    attachments.append(file_info.filename)
                    if file_info.filename.endswith('.txt'):
                        try:
                            content = z.read(file_info.filename).decode('utf-8')
                            text_blocks.append(content)
                        except Exception as decode_e:
                            logger.warning(f"Could not decode text file in zip: {decode_e}")
                            
            return self._standard_response(
                parser_name="zip",
                success=True,
                text="\\n".join(text_blocks),
                attachments=attachments,
                confidence=0.8
            )
        except Exception as e:
            logger.error(f"ZIPParser failed: {e}")
            return self._standard_response(parser_name="zip", success=False, confidence=0.0)
"""
with open(os.path.join(base_path, "zip_parser.py"), "w") as f:
    f.write(zip_parser_code)


# 6. docx_parser.py
docx_parser_code = """import logging
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class DOCXParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            import docx
            doc = docx.Document(source)
            text_blocks = [para.text for para in doc.paragraphs if para.text.strip()]
            
            tables = []
            for table in doc.tables:
                grid = []
                for row in table.rows:
                    grid.append([cell.text.strip() for cell in row.cells])
                tables.append(grid)
                
            return self._standard_response(
                parser_name="docx",
                success=True,
                text="\\n".join(text_blocks),
                tables=tables,
                confidence=0.95
            )
        except ImportError:
            logger.error("python-docx not installed.")
            return self._standard_response(parser_name="docx", success=False, confidence=0.0)
        except Exception as e:
            logger.error(f"DOCXParser failed: {e}")
            return self._standard_response(parser_name="docx", success=False, confidence=0.0)
"""
with open(os.path.join(base_path, "docx_parser.py"), "w") as f:
    f.write(docx_parser_code)


# 7. txt_parser.py
txt_parser_code = """import logging
from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)

class TXTParser(BaseParser):
    def parse(self, source: str, **kwargs) -> dict:
        try:
            with open(source, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                
            return self._standard_response(
                parser_name="txt",
                success=True,
                text=content,
                confidence=0.9
            )
        except Exception as e:
            logger.error(f"TXTParser failed: {e}")
            return self._standard_response(parser_name="txt", success=False, confidence=0.0)
"""
with open(os.path.join(base_path, "txt_parser.py"), "w") as f:
    f.write(txt_parser_code)


# Now we must update pdf_parser.py and html_parser.py to inherit from BaseParser so they are compatible
print("Finished rewriting basic parsers.")
